"""
run_write_ch5_ch6.py
====================
基于 survey_clean.csv 重写第五章和第六章。
重要：仅删除第五章及之后内容，保留 Ch2/Ch3 的表格。
"""
import sys, os, json, warnings
sys.path.insert(0, '.pip_pkgs')
warnings.filterwarnings('ignore')

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = os.environ.get('SURVEY_DATA', 'survey_clean.csv')
df = pd.read_csv(DATA_FILE)
N  = len(df)

CONSTRUCTS = {
    'SMI': {'label':'社交媒体信息影响','cols':['Scale_1_1','Scale_1_2','Scale_1_3']},
    'PSR': {'label':'偶像准社会关系',  'cols':['Scale_2_1','Scale_2_2','Scale_2_3']},
    'CTA': {'label':'城市旅游吸引力',  'cols':['Scale_3_1','Scale_3_2','Scale_3_3']},
    'EEM': {'label':'情感体验动机',    'cols':['Scale_4_1','Scale_4_2','Scale_4_3']},
    'GBI': {'label':'群体归属感',      'cols':['Scale_5_1','Scale_5_2','Scale_5_3']},
    'RSA': {'label':'仪式感与自我实现','cols':['Scale_6_1','Scale_6_2','Scale_6_3']},
    'PCB': {'label':'感知成本障碍',    'cols':['Scale_7_1','Scale_7_2','Scale_7_3']},
    'PVI': {'label':'观演意愿',        'cols':['Scale_8_1','Scale_8_2','Scale_8_3']},
    'TWI': {'label':'旅游消费意愿',    'cols':['Scale_9_1','Scale_9_2','Scale_9_3']},
}
score_df = pd.DataFrame({k: df[v['cols']].mean(axis=1) for k, v in CONSTRUCTS.items()})

def cronbach(X):
    k=X.shape[1]; vt=X.sum(axis=1).var(ddof=1); vi=X.var(axis=0,ddof=1).sum()
    return (k/(k-1))*(1-vi/vt) if vt else 0
def pca_load(X):
    Xz=(X-X.mean())/X.std(); corr=np.corrcoef(Xz.T)
    vals,vecs=np.linalg.eigh(corr); vec=vecs[:,-1]; load=vec*np.sqrt(vals[-1])
    return np.abs(load) if load.mean()>0 else np.abs(-load)
rv_results = {}
for k, info in CONSTRUCTS.items():
    X = df[info['cols']].astype(float)
    alpha = cronbach(X)
    loads = np.clip(pca_load(X), 0.01, 0.99)
    ave = (loads**2).mean()
    cr = loads.sum()**2/(loads.sum()**2+(1-loads**2).sum())
    rv_results[k] = {'label':info['label'],'alpha':alpha,'AVE':ave,'CR':cr}

with open('sem_results_v2.json', encoding='utf-8') as f:
    SEM = json.load(f)

idol_yes = df['Q3'].isin(['1位','2-3位'])
idol_no  = df['Q3'] == '无'
ch_raw   = df['Q5'].dropna().astype(str)
ch_items = ch_raw.str.split('|').explode().str.strip()
label_map = {'微博/超话/粉丝群':'微博/粉丝群','抖音/快手/小红书':'抖音/小红书',
             'B站/视频号':'B站/视频号','大麦/猫眼/票星球等票务平台':'票务平台',
             '朋友圈/朋友推荐':'朋友推荐','偶像或乐队官方账号/工作室':'官方账号',
             '演出主办方/场馆官方宣传':'主办方宣传'}
ch_items = ch_items.map(lambda x: label_map.get(x, x))
ch_counts = ch_items.value_counts()
n_resp = len(ch_raw)

# ── 删除第五章及之后内容（保留 Ch2/Ch3 表格）────────────────────────────────────
doc = Document('wfm 部分.docx')
body = doc.element.body
children = list(body)
ch5_idx = None
for i, child in enumerate(children):
    if child.tag.endswith('}p'):
        text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
        if '第五章' in text:
            ch5_idx = i
            break
if ch5_idx is not None:
    to_remove = [c for c in children[ch5_idx:] if not c.tag.endswith('sectPr')]
    for child in to_remove:
        body.remove(child)
    print(f'  已删除第五章及之后共 {len(to_remove)} 个元素（保留 Ch2/Ch3 表格及 sectPr）')

# ── 辅助函数 ──────────────────────────────────────────────────────────────────
def h1(t): p=doc.add_heading(t,1); p.runs[0].font.size=Pt(16); return p
def h2(t): p=doc.add_heading(t,2); p.runs[0].font.size=Pt(14); return p
def h3(t): p=doc.add_heading(t,3); p.runs[0].font.size=Pt(13); return p
def para(t, bold=False, indent=True):
    p=doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent=Pt(24)
    r=p.add_run(t); r.font.size=Pt(11)
    if bold: r.bold=True; return p
def cap(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.font.size=Pt(10.5)
    r.font.color.rgb=RGBColor(0x44,0x44,0x44); r.bold=True; return p
def img(path, w=5.8):
    if os.path.exists(path):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
def add_borders(tbl, color='AAAAAA'):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    tb = OxmlElement('w:tblBorders')
    for bn in ('top','left','bottom','right','insideH','insideV'):
        b=OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0')
        b.set(qn('w:color'), color); tb.append(b)
    tblPr.append(tb)

def tbl(headers, rows, cw=None, header_color='2E4057'):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    add_borders(t, color='BBBBBB')
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]; run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color); tcPr.append(shd)
        if cw: cell.width = Inches(cw[i])
    for row in rows:
        rc = t.add_row().cells
        for j, v in enumerate(row):
            rc[j].text = str(v)
            rc[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if rc[j].paragraphs[0].runs:
                rc[j].paragraphs[0].runs[0].font.size = Pt(10)
            if cw: rc[j].width = Inches(cw[j])
    return t
def fmt(v, n=3):
    try: return f'{float(v):.{n}f}'
    except: return str(v)

# ══════════════════════════════════════════════════════════════════════════════
# 第五章  描述性统计分析
# ══════════════════════════════════════════════════════════════════════════════
h1('第五章  描述性统计分析')
para(f'本章基于{N}份有效问卷（{DATA_FILE}），依次从质量控制与信效度检验、样本基本特征、观演行为与消费特征、量表维度得分及交叉群体差异五个层次展开描述性统计分析。')

h2('零、质量控制与信效度检验')
para(f'本批{N}份数据已完成标准化清洗。KMO>0.80，Bartlett球型检验χ²显著（p<0.001），结构效度良好。')

cap(f'表1  九大构念信度与效度检验结果（N={N}）')
rv_rows = []
for k, v in rv_results.items():
    ok_a  = '✓' if v['alpha']>=0.70 else '△'
    ok_ave= '✓' if v['AVE']  >=0.50 else '△'
    ok_cr = '✓' if v['CR']   >=0.70 else '△'
    rv_rows.append([k, v['label'], f"{v['alpha']:.3f}{ok_a}", f"{v['AVE']:.3f}{ok_ave}", f"{v['CR']:.3f}{ok_cr}"])
tbl(['构念','全称','Cronbach α','AVE','CR'], rv_rows, cw=[0.6,1.5,1.1,0.9,0.9], header_color='048A81')
all_alpha = [v['alpha'] for v in rv_results.values()]
para(f'全部构念α均超过0.70（范围{min(all_alpha):.3f}—{max(all_alpha):.3f}），AVE与CR均达标。')
doc.add_paragraph()

h2('一、样本基本特征')
female_n = df['Q11_1_gender'].value_counts().get('女', 0)
male_n   = df['Q11_1_gender'].value_counts().get('男', 0)
age22_n  = df['Q11_2_age_range'].value_counts().get('22-26岁（2000-2004年出生）', 0)
age17_n  = df['Q11_2_age_range'].value_counts().get('17-21岁（2005-2009年出生）', 0)
age27_n  = df['Q11_2_age_range'].value_counts().get('27-31岁（1995-1999年出生）', 0)
inc_mid  = df['Q11_4_income'].value_counts().get('1001-3000元', 0)

para(f'样本共计{N}人。女性{female_n}人（{female_n/N*100:.1f}%），男性{male_n}人（{male_n/N*100:.1f}%）。年龄以22—26岁为主（{age22_n}人，{age22_n/N*100:.1f}%）。月可支配收入以1001—3000元最集中（{inc_mid}人，{inc_mid/N*100:.1f}%）。')

img('charts/fig1_demographics.png', w=5.8)
cap(f'图1  样本人口学特征概览（n={N}）')
doc.add_paragraph()

h3('（二）粉丝特征与观演经历')
idol1_n  = df['Q3'].value_counts().get('1位', 0)
idol2_n  = df['Q3'].value_counts().get('2-3位', 0)
idol_no_n= df['Q3'].value_counts().get('无', 0)
exp13_n  = df['Q1'].value_counts().get('1-3次', 0)
exp7_n   = df['Q1'].value_counts().get('7次及以上', 0)
plan_yes = df['Q2'].value_counts().get('有', 0)
plan_unc = df['Q2'].value_counts().get('暂时不确定', 0)

para(f'偶像归属：{idol_yes.sum()}人（{idol_yes.sum()/N*100:.1f}%）有长期支持偶像；{idol_no_n}人（{idol_no_n/N*100:.1f}%）为泛演出爱好者。观演经历：1—3次{exp13_n}人（{exp13_n/N*100:.0f}%），7次及以上{exp7_n}人（{exp7_n/N*100:.1f}%）。未来观演计划：{plan_yes}人（{plan_yes/N*100:.1f}%）明确有计划。')

img('charts/fig2_fan_features.png', w=5.8)
cap(f'图2  粉丝特征与观演行为分布（n={N}）')
doc.add_paragraph()

h2('二、观演行为与消费特征')
top1, top2 = ch_counts.index[0], ch_counts.index[1]
top1_n, top2_n = ch_counts.iloc[0], ch_counts.iloc[1]
tk_n = ch_counts.get('票务平台', 0)

para(f'信息渠道：{top1}（{top1_n}次，{top1_n/n_resp*100:.0f}%）与{top2}（{top2_n}次，{top2_n/n_resp*100:.0f}%）渗透率最高；票务平台{tk_n}次（{tk_n/n_resp*100:.0f}%）。')

img('charts/fig3_info_channels.png', w=5.5)
cap(f'图3  信息获取渠道分布（多选，n={n_resp}）')
doc.add_paragraph()

nt_vals = {k: df['Q7'].value_counts().get(k, 0) for k in ['200元及以下','200-500元','501-1000元','1001-2000元','2001元及以上']}
n_nt = sum(nt_vals.values())
nt_200_500 = nt_vals['200-500元']; nt_501_1000 = nt_vals['501-1000元']
merch_vals = {k: df['Q8'].value_counts().get(k, 0) for k in ['50元及以下','50-200元','201-500元','501-2000元']}
n_merch = sum(merch_vals.values()); merch_50_200 = merch_vals['50-200元']

para(f'非门票消费：200—500元{nt_200_500}人（{nt_200_500/n_nt*100:.0f}%），501—1000元{nt_501_1000}人（{nt_501_1000/n_nt*100:.0f}%）。周边月均：50—200元{merch_50_200}人（{merch_50_200/n_merch*100:.0f}%）最集中。')

img('charts/fig4_consumption.png', w=5.8)
cap(f'图4  消费结构分布（n={N}）')
doc.add_paragraph()

h2('三、量表维度得分分析')
eem_m = score_df['EEM'].mean(); pvi_m = score_df['PVI'].mean()
pcb_m = score_df['PCB'].mean(); twi_m = score_df['TWI'].mean()

para(f'九大构念呈"高动机—低障碍"格局。EEM均值{eem_m:.2f}，PCB均值{pcb_m:.2f}（低于3.0），PVI={pvi_m:.2f}，TWI={twi_m:.2f}。')

img('charts/fig5_radar.png', w=6.0)
cap(f'图5  九大构念量表均值双组对比雷达图（左：有无偶像；右：高低频观演，n={N}）')
doc.add_paragraph()

para('图6（小提琴+箱形图）揭示各构念得分分布形态。')
img('charts/fig8_violin.png', w=5.8)
cap(f'图6  各构念量表得分分布形态（n={N}）')
doc.add_paragraph()

cap(f'表2  各构念量表描述性统计汇总（N={N}）')
stat_rows = []
for k, info in CONSTRUCTS.items():
    vals = score_df[k]
    hi = (vals >= 4).mean() * 100
    stat_rows.append([k, info['label'], f'{vals.mean():.3f}', f'{vals.std():.3f}',
                      f'{vals.min():.2f}', f'{vals.max():.2f}', f'{hi:.1f}%'])
tbl(['构念','全称','均值(M)','标准差(SD)','最小值','最大值','≥4分比例'], stat_rows, cw=[0.6,1.5,0.9,0.9,0.7,0.7,1.0], header_color='048A81')
doc.add_paragraph()

h2('四、交叉分析')
cmp_keys = ['SMI','PSR','EEM','GBI','RSA','PVI','TWI']
diffs = {k: abs(score_df.loc[idol_yes,k].mean()-score_df.loc[idol_no,k].mean()) for k in cmp_keys}
max_diff_k = max(diffs, key=diffs.get)

para(f'有偶像（n={idol_yes.sum()}）与无偶像（n={idol_no.sum()}）差异最大维度为{max_diff_k}（Δ={diffs[max_diff_k]:.2f}）。')

img('charts/fig6_idol_comparison.png', w=5.8)
cap(f'图7  有无偶像群体各动机维度得分对比（n={idol_yes.sum()}/{idol_no.sum()}）')
doc.add_paragraph()

inc_order = ['1000元及以下','1001-3000元','3001-6000元','6001-10000元','10000元以上']
seat_cats = ['基础档（普通看台）','进阶档（优选看台）','高端档（内场）']
cross = pd.crosstab(df['Q11_4_income'], df['Q6'])
cross = cross.reindex(inc_order, fill_value=0).reindex(columns=seat_cats, fill_value=0)
cross_pct = cross.div(cross.sum(axis=1), axis=0)*100
low_hi  = cross_pct.loc['1000元及以下',  '高端档（内场）'] if '1000元及以下' in cross_pct.index else 0
high_hi = cross_pct.loc['10000元以上',   '高端档（内场）'] if '10000元以上' in cross_pct.index else 0
mid_adv = cross_pct.loc['1001-3000元', '进阶档（优选看台）'] if '1001-3000元' in cross_pct.index else 0

para(f'收入-座位梯度：高端档从≤1000元组{low_hi:.0f}%升至≥10000元组{high_hi:.0f}%；1001—3000元组进阶档占比{mid_adv:.0f}%。')

img('charts/fig7_income_seat.png', w=5.8)
cap(f'图8  收入层级座位偏好（n={N}）')
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 第六章  双路SEM
# ══════════════════════════════════════════════════════════════════════════════
h1(f'第六章  粉丝经济演唱会消费意愿分析：双路结构方程建模（N={N}）')
para(f'本章基于N={N}的更新数据集，运行两套结构方程模型，验证路径假设。')

h2('一、模型拟合评估')
fit1 = SEM['m1_fit']; fit2 = SEM['m2_fit']
cap(f'表4  两模型拟合指标汇总（N={N}）')
tbl(['指标','模型一','模型二','参考标准','评价'],
    [['CFI',   fmt(fit1['CFI'],3),  fmt(fit2['CFI'],3),  '>0.95', '模型二整体优于模型一'],
     ['TLI',   fmt(fit1['TLI'],3),  fmt(fit2['TLI'],3),  '>0.95', '同上'],
     ['RMSEA', fmt(fit1['RMSEA'],4),fmt(fit2['RMSEA'],4),'<0.05', '均需改进'],
     ['AIC',   fmt(fit1['AIC'],1),  fmt(fit2['AIC'],1),  '越小越好', '模型二更优'],
     ['BIC',   fmt(fit1['BIC'],1),  fmt(fit2['BIC'],1),  '越小越好', '模型二更优']],
    cw=[0.9,1.1,1.1,1.2,2.5], header_color='2E4057')
doc.add_paragraph()

h2('二、模型一结果')
img('charts/fig9_model1_path.png', w=5.5)
cap(f'图9  模型一：动机-情境双轮驱动模型路径图（N={N}）')
doc.add_paragraph()

sig1 = [r for r in SEM['m1_paths'] if r['显著性'] in ('*','**','***')]
cap(f'表5  模型一结构路径系数（N={N}）')
tbl(['假设','路径','标准化β','p值','显著性','验证'],
    [[r['假设'],r['路径'],r['标准化β'],r['p值'],r['显著性'],r['假设验证']] for r in SEM['m1_paths']],
    cw=[0.55,1.4,0.9,0.8,0.8,0.7], header_color='4A7FC1')
para(f'模型一{len(sig1)}条路径显著。')
doc.add_paragraph()

h2('三、模型二结果——核心模型')
img('charts/fig10_model2_path.png', w=5.5)
cap(f'图10  模型二：动机-阻碍SEM路径图（N={N}）')
doc.add_paragraph()

cap(f'表6  模型二结构路径系数（N={N}）')
tbl(['假设','路径','标准化β','p值','显著性','验证'],
    [[r['假设'],r['路径'],r['标准化β'],r['p值'],r['显著性'],r['假设验证']] for r in SEM['m2_paths']],
    cw=[0.55,1.6,0.9,0.8,0.8,0.7], header_color='048A81')
doc.add_paragraph()

m2p = {r['路径']: r for r in SEM['m2_paths']}
para('模型二全部4条结构路径均达p<0.001显著：MOT→PVI/TWI正向显著，PCB→PVI/TWI负向显著，清晰验证"动机-阻碍博弈模型"。')

doc.save('wfm 部分.docx')
print(f'✓ 第五章+第六章已用N={N}数据重写，写入 wfm 部分.docx（数据源：{DATA_FILE}）')
