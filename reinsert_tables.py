# -*- coding: utf-8 -*-
"""
重新插入被意外删除的KMO表格和Fornell-Larcker矩阵
锚点：para [140] 的 Kaiser-Meyer-Olkin 段落
插入位置：该段落之后，"一、样本基本特征" 之前
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document("wfm \u90e8\u5206.docx")
body = doc.element.body

def get_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t")))

def make_para(text, bold=False, indent=False, fs=10.5, align=None):
    p = doc.add_paragraph()
    p._element.getparent().remove(p._element)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(fs)
        r.bold = bold
    return p._element

def add_border(tbl_obj, color="999999"):
    tblPr = tbl_obj._tbl.tblPr
    tb = OxmlElement("w:tblBorders")
    for bn in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{bn}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), color)
        tb.append(b)
    tblPr.append(tb)

def set_shade(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)

def make_table(headers, rows, cw, hdr_color="1A5276", fs=9.5, bold_c0=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t._element.getparent().remove(t._element)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_border(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h; c.width = Inches(cw[i])
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]; run.bold = True; run.font.size = Pt(fs)
        run.font.color.rgb = RGBColor(255, 255, 255); set_shade(c, hdr_color)
    for ri, row in enumerate(rows):
        rc = t.add_row().cells
        bg = "FFFFFF" if ri % 2 == 0 else "F2F4F7"
        for j, v in enumerate(row):
            rc[j].text = str(v); rc[j].width = Inches(cw[j])
            p = rc[j].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(fs)
                if bold_c0 and j == 0: p.runs[0].bold = True
            set_shade(rc[j], bg)
    return t._element

# ── Find anchor: Kaiser-Meyer-Olkin paragraph ────────────────────────────────
anchor_el = None
for el in list(body):
    if el.tag.endswith("}p") and "Kaiser-Meyer-Olkin" in get_text(el):
        anchor_el = el
        break

if anchor_el is None:
    print("ERROR: Could not find KMO anchor paragraph")
    exit(1)

print(f"Found anchor: {get_text(anchor_el)[:80]}")

# ── Build all content to insert ──────────────────────────────────────────────
new_els = []

# 1. KMO table caption
new_els.append(make_para(
    "\u88681a\u00a0\u00a0KMO\u53d6\u6837\u9002\u5207\u6027\u68c0\u9a8c\u4e0eBartlett\u7403\u578b\u68c0\u9a8c\u7ed3\u679c\uff08N=712\uff0c27\u9898\uff09",
    fs=10.5, align=WD_ALIGN_PARAGRAPH.CENTER
))

# 2. KMO table
kmo_headers = ["\u68c0\u9a8c\u9879\u76ee", "\u7edf\u8ba1\u91cf", "\u6570\u5024", "\u5224\u65ad\u6807\u51c6", "\u7ed3\u8bba"]
kmo_rows = [
    ["KMO\u53d6\u6837\u9002\u5207\u6027", "KMO\u5024", "0.948",
     "\u22650.70\uff08\u5408\u683c\uff09\uff1b\u22650.90\uff08\u4f18\u79c0\uff09", "\u4f18\u79c0\uff0c\u9ad8\u5ea6\u9002\u5408\u56e0\u5b50\u5206\u6790"],
    ["Bartlett\u7403\u578b\u68c0\u9a8c", "\u8fd1\u4f3c\u03c7\u00b2", "34649.61",
     "\u663e\u8457\u6027p<0.05", "\u76f8\u5173\u77e9\u9635\u975e\u5355\u4f4d\u77e9\u9635"],
    ["", "\u81ea\u7531\u5ea6df", "351", "\u2014", "\u2014"],
    ["", "\u663e\u8457\u6027p\u5024", "<0.001", "p<0.001", "\u2713 \u5f3a\u70c8\u62d2\u7edd\u96f6\u5047\u8bbe"],
]
new_els.append(make_table(kmo_headers, kmo_rows, [1.5, 1.4, 1.0, 2.0, 2.1]))

# 3. KMO result text
new_els.append(make_para(
    "KMO\u5024=0.948\uff0c\u5c5e\u4e8e\u300c\u4f18\u79c0\u300d\u7ea7\u522b\uff08Kaiser, 1974\uff09\uff0c\u8868\u660e\u6837\u672c\u6570\u636e\u9ad8\u5ea6\u9002\u5408\u8fdb\u884c\u56e0\u5b50\u5206\u6790\uff1b"
    "Bartlett\u7403\u578b\u68c0\u9a8c\u03c7\u00b2(351)=34649.61\uff0cp<0.001\uff0c\u5f3a\u70c8\u62d2\u7edd"
    "\u300c\u76f8\u5173\u77e9\u9635\u4e3a\u5355\u4f4d\u77e9\u9635\u300d\u7684\u96f6\u5047\u8bbe\uff0c\u786e\u8ba4\u53d8\u91cf\u95f4\u5b58\u5728\u5145\u5206\u7684\u5171\u540c\u56e0\u5b50\u7ed3\u6784\uff0c"
    "\u4e3a\u540e\u7eedSEM\u7684\u5efa\u6784\u63d0\u4f9b\u4e86\u4e25\u683c\u7684\u7edf\u8ba1\u524d\u63d0\u3002",
    indent=True, fs=10.5
))

# 4. Discriminant validity heading
new_els.append(make_para(
    "\uff08\u4e8c\uff09\u533a\u5206\u6548\u5ea6\u68c0\u9a8c\u2014\u2014Fornell-Larcker\u51c6\u5219",
    bold=True, fs=10.5
))
new_els.append(make_para(
    "\u4f9d\u636eFornell & Larcker\uff081981\uff09\u51c6\u5219\uff0c\u5f53\u67d0\u4e00\u6784\u5ff5\u7684AVE\u5e73\u65b9\u6839\uff08\u221aAVE\uff09\u5927\u4e8e\u8be5\u6784\u5ff5\u4e0e\u6240\u6709\u5176\u4ed6\u6784\u5ff5\u7684\u76f8\u5173\u7cfb\u6570\u65f6\uff0c"
    "\u5373\u53ef\u8ba4\u5b9a\u533a\u5206\u6548\u5ea6\u8fbe\u6807\u3002\u4ee5\u4e0b\u5206\u522b\u62a5\u544a\u65b9\u6848\u4e00\uff088\u6784\u5ff5\uff09\u4e0e\u65b9\u6848\u4e8c\uff086\u6784\u5ff5\uff09\u7684Fornell-Larcker\u77e9\u9635\u3002",
    indent=True, fs=10.5
))

# 5. FL Table 1b caption
new_els.append(make_para(
    "\u88681b\u00a0\u00a0\u65b9\u6848\u4e00\u533a\u5206\u6548\u5ea6\u77e9\u9635\uff08Fornell-Larcker\u51c6\u5219\uff0cn=712\uff09"
    "\n\uff08\u5bf9\u89d2\u7ebf\u4e3a\u5404\u6784\u5ff5\u221aAVE\uff0c\u4e0b\u4e09\u89d2\u4e3a\u4e24\u4e24\u6784\u5ff5\u95f4\u76f8\u5173\u7cfb\u6570\uff09",
    fs=10.5, align=WD_ALIGN_PARAGRAPH.CENTER
))

# 6. FL Table 1b data
fl1_h = ["\u6784\u5ff5", "SMI", "PSR", "CTA", "EEM", "GBI", "RSA", "PVI", "TWI"]
fl1_r = [
    ["SMI",  "[0.915]", "",       "",       "",       "",       "",       "",       ""],
    ["PSR",  "0.944",  "[0.976]", "",       "",       "",       "",       "",       ""],
    ["CTA",  "0.116",  "0.012",  "[0.905]", "",       "",       "",       "",       ""],
    ["EEM",  "0.464",  "0.451",  "0.105",  "[0.865]", "",       "",       "",       ""],
    ["GBI",  "0.867",  "0.833",  "0.118",  "0.499",  "[0.924]", "",       "",       ""],
    ["RSA",  "0.913",  "0.873",  "0.141",  "0.391",  "0.896",  "[0.937]", "",       ""],
    ["PVI",  "0.885",  "0.858",  "0.255",  "0.551",  "0.833",  "0.835",  "[0.941]", ""],
    ["TWI",  "0.600",  "0.533",  "0.535",  "0.315",  "0.521",  "0.563",  "0.663",  "[0.816]"],
]
new_els.append(make_table(fl1_h, fl1_r, [0.72]*9, bold_c0=True, fs=9.0))

# 7. FL1 note
new_els.append(make_para(
    "\u6ce8\uff1a\u65b9\u6848\u4e00\u4e2d\uff0cSMI-PSR\u3001SMI-GBI\u3001SMI-RSA\u3001PSR-GBI\u3001PSR-RSA\u3001GBI-RSA\u7b49\u6784\u5ff5\u5bf9\u76f8\u5173\u7cfb\u6570"
    "\u8d85\u8fc7\u5404\u81ea\u221aAVE\uff0c\u63d0\u793a\u8fd9\u4e9b\u5185\u56e0\u52a8\u673a\u53d8\u91cf\u5728\u91cf\u8868\u5c42\u9762\u5b58\u5728\u8f83\u9ad8\u5171\u7ebf\u6027\uff0c"
    "\u662f\u65b9\u6848\u4e00\u6700\u7ec8\u7cbe\u7b80\u4e3a\u65b9\u6848\u4e8c\uff08\u5c06EEM/GBI/RSA\u805a\u5408\u4e3aMOT\uff09\u7684\u7edf\u8ba1\u4f9d\u636e\u3002",
    indent=True, fs=9.5
))

# 8. FL Table 1c caption
new_els.append(make_para(
    "\u88681c\u00a0\u00a0\u65b9\u6848\u4e8c\u533a\u5206\u6548\u5ea6\u77e9\u9635\uff08Fornell-Larcker\u51c6\u5219\uff0cn=712\uff09"
    "\n\uff08\u5bf9\u89d2\u7ebf\u4e3a\u5404\u6784\u5ff5\u221aAVE\uff0c\u4e0b\u4e09\u89d2\u4e3a\u4e24\u4e24\u6784\u5ff5\u95f4\u76f8\u5173\u7cfb\u6570\uff09",
    fs=10.5, align=WD_ALIGN_PARAGRAPH.CENTER
))

# 9. FL Table 1c data
fl2_h = ["\u6784\u5ff5", "EEM", "GBI", "RSA", "PCB", "PVI", "TWI"]
fl2_r = [
    ["EEM",  "[0.865]", "",        "",        "",        "",        ""],
    ["GBI",  "0.499",  "[0.924]",  "",        "",        "",        ""],
    ["RSA",  "0.391",  "0.896",   "[0.937]",  "",        "",        ""],
    ["PCB",  "-0.276", "-0.488",  "-0.536",  "[0.897]",  "",        ""],
    ["PVI",  "0.551",  "0.833",   "0.835",   "-0.645",  "[0.941]",  ""],
    ["TWI",  "0.315",  "0.521",   "0.563",   "-0.804",  "0.663",   "[0.816]"],
]
new_els.append(make_table(fl2_h, fl2_r, [0.72, 0.90, 0.90, 0.90, 0.90, 0.90, 0.90],
                           bold_c0=True, fs=9.5))

# 10. FL2 note
new_els.append(make_para(
    "\u6ce8\uff1a\u65b9\u6848\u4e8c\u4e2d\uff0cPCB\u4e0eTWI\u95f4\u76f8\u5173\u7cfb\u6570(-0.804)\u7edd\u5bf9\u5024\u8d85\u8fc7\u5404\u81ea\u221aAVE\uff0c"
    "\u5c5e\u4e8e\u7406\u8bba\u5408\u7406\u8303\u56f4\u5185\u7684\u5f3a\u8d1f\u5411\u5224\u522b\u6027\u5173\u8054\uff08\u800c\u975e\u6d4b\u91cf\u6df7\u6dc6\uff09\uff0c"
    "\u4e0e\u6a21\u578b\u4e8c\u4e2dPCB\u2192TWI\u8def\u5f84\u7cfb\u6570(\u03b2=-0.459, p<0.001)\u9ad8\u5ea6\u663e\u8457\u76f8\u4e92\u5370\u8bc1\u3002",
    indent=True, fs=9.5
))

# ── Insert all elements after anchor_el ──────────────────────────────────────
for el in reversed(new_els):
    anchor_el.addnext(el)

doc.save("wfm \u90e8\u5206.docx")
print(f"Done! Inserted {len(new_els)} elements after KMO anchor paragraph.")

# Verify
doc2 = Document("wfm \u90e8\u5206.docx")
body2 = doc2.element.body
def gt(el): return "".join(t.text or "" for t in el.iter(qn("w:t")))
checks = ["34649.61", "Fornell-Larcker", "\u65b9\u6848\u4e00\u533a\u5206\u6548\u5ea6", "\u65b9\u6848\u4e8c\u533a\u5206\u6548\u5ea6"]
for kw in checks:
    found = any(kw in gt(el) for el in list(body2) if el.tag.endswith("}p") or el.tag.endswith("}tbl"))
    print(f"  {'OK' if found else 'MISSING'}: {kw}")
