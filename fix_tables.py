"""
fix_tables.py
=============
1. 在 Ch2/Ch3 缺失表格标题后插入对应表格
2. 更新 docx 中的雷达图(fig5) 为双组版本
3. 更新 docx 中图5 文字说明
"""
import sys, os
sys.path.insert(0, '.pip_pkgs')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document('wfm 部分.docx')

# ── 辅助函数 ──────────────────────────────────────────────────────────────────
def add_borders(tbl, color='BBBBBB'):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    tb = OxmlElement('w:tblBorders')
    for bn in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
        tb.append(b)
    tblPr.append(tb)

def make_table(headers, rows, header_shade='4A90D9', cw=None):
    """Create a table with colored header."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(t)
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        # Cell shading
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_shade); tcPr.append(shd)
        if cw: cell.width = Inches(cw[i])
    # Data rows
    for row_data in rows:
        r = t.add_row()
        for j, val in enumerate(row_data):
            cell = r.cells[j]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].font.size = Pt(9.5)
            if cw: cell.width = Inches(cw[j])
    return t

def insert_tbl_after_para(para, tbl_obj):
    """Insert table element right after a paragraph element."""
    para._element.addnext(tbl_obj._tbl)

def find_para(text_contains):
    """Find paragraphs containing given text."""
    return [p for p in doc.paragraphs if text_contains in p.text]

# ══════════════════════════════════════════════════════════════════════════════
# 表1  混合研究法技术路线
# ══════════════════════════════════════════════════════════════════════════════
t1_headers = ['研究阶段', '主要工作', '研究方法', '分析工具', '核心产出']
t1_rows = [
    ['① 文献综述\n与概念建模', '系统梳理粉丝经济、Z世代消费、\n演唱会经济三大文献领域', '文献分析法、\n概念整合', 'Zotero/知网/\nWoS', '10个初始潜变量\n及概念框架'],
    ['② 定性研究\n与访谈', '现场拦截访谈（n=17）、深度访谈\n（n=6）、焦点小组（2场）;\nLLM文本分析（社媒评论）', '拦截访谈、IDI、\nFGD、LLM', '词云图/扎根理论/\nLLM（另行报告）', '校正量表题项池\n痛点/动机清单'],
    ['③ 预调查\n与量表筛选', '小样本定向投放（n≈50），\n依次进行CITC、Cronbach α、\nEFA三轮统计筛选', 'CITC法、信度\n检验、EFA', 'Python/\n(pingouin/sklearn)', '精简量表（9个构念）\n进入正式问卷'],
    ['④ 大规模\n定量调查', '分层多阶段抽样问卷，\n线上+线下双渠道正式发放，\n回收清洗后有效样本300份', '分层整群抽样、\n便利/配额抽样', '问卷星/\n线下纸质', '有效问卷 N=300\n（清洗后最终样本）'],
    ['⑤ 多模型\n综合分析', '双路SEM、DCM离散选择模型、\n线性加权多目标规划、\n用户聚类画像及前端展示', 'SEM、DCM、\n聚类、规划', 'Python\n(semopy/sklearn/\ndocplex)', '研究结论与\n决策建议'],
]
paras_t1 = find_para('混合研究法技术路线')
if paras_t1:
    t1 = make_table(t1_headers, t1_rows, header_shade='2E4057', cw=[0.9,2.3,1.2,1.2,1.5])
    insert_tbl_after_para(paras_t1[0], t1)
    print('✓ 表1 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表2  10个初始潜变量一览表
# ══════════════════════════════════════════════════════════════════════════════
t2_headers = ['编号', '代码', '构念中文名称', '维度层次', '初始\n题项数', '理论来源/定性方法确认']
t2_rows = [
    ['1',  'SMI', '社交媒体信息影响',   '外因-背景情境', '3', '社会影响理论（SIT）；访谈高频词"刷到推送"'],
    ['2',  'PSR', '偶像准社会关系',     '外因-背景情境', '3', '准社会互动理论（PSI）；IDI核心主题之一'],
    ['3',  'CTA', '城市旅游吸引力',     '外因-背景情境', '3', '城市旅游动机文献；拦访中"天津来玩顺便看"'],
    ['4',  'EEM', '情感体验动机',       '内因-动机层',   '3', '体验经济理论（Pine & Gilmore）；FGD核心主题'],
    ['5',  'GBI', '群体归属感',         '内因-动机层',   '3', '社会认同理论；FGD关键词"和朋友一起追星"'],
    ['6',  'RSA', '仪式感与自我实现',   '内因-动机层',   '3', '仪式理论（Durkheim）+ Maslow需求层次'],
    ['7',  'PCB', '感知成本障碍',       '内因-阻碍层',   '3', '价值-障碍框架；拦访最高频痛点'],
    ['8',  'SC',  '服务体验顾虑',       '内因-阻碍层',   '3', '服务质量理论；拦访提及"排队太久/安保混乱"'],
    ['9',  'PVI', '观演意愿',           '因变量-行为意向', '3', '计划行为理论（TPB）'],
    ['10', 'TWI', '旅游/消费延伸意愿',  '因变量-行为意向', '3', '旅游意向文献；"演唱会带动城市消费"议题'],
]
paras_t2 = find_para('初始潜变量一览表')
if paras_t2:
    t2 = make_table(t2_headers, t2_rows, header_shade='048A81', cw=[0.4,0.55,1.35,1.15,0.5,2.1])
    insert_tbl_after_para(paras_t2[0], t2)
    print('✓ 表2 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表3  Cronbach's α信度判断标准
# ══════════════════════════════════════════════════════════════════════════════
t3_headers = ['Cronbach\'s α 范围', '信度水平', '预调查处理方式']
t3_rows = [
    ['α ≥ 0.90', '优秀（Excellent）', '保留全部题项，量表质量高'],
    ['0.80 ≤ α < 0.90', '良好（Good）', '保留，建议微调题项措辞'],
    ['0.70 ≤ α < 0.80', '可接受（Acceptable）', '保留，结合CITC优化'],
    ['0.60 ≤ α < 0.70', '存疑（Questionable）', '删除CITC最低题项后重检'],
    ['α < 0.60', '不可接受（Poor）', '该构念整体剔除或重设题项'],
]
paras_t3 = find_para('Cronbach\'s α信度判断标准')
if not paras_t3:
    paras_t3 = find_para('Cronbach')
# find the specific table caption
all_cronbach = [p for p in doc.paragraphs if 'α信度判断标准' in p.text or 'Cronbach\'s α信度判断标准' in p.text]
if all_cronbach:
    t3 = make_table(t3_headers, t3_rows, header_shade='6B4E71', cw=[1.5,1.5,3.1])
    insert_tbl_after_para(all_cronbach[0], t3)
    print('✓ 表3 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表4  EFA题项与变量筛选标准
# ══════════════════════════════════════════════════════════════════════════════
t4_headers = ['检验指标', '筛选标准（阈值）', '未达标处理方式', '备注']
t4_rows = [
    ['因子载荷（λ）', '≥ 0.50', '删除该题项', '与主因子的相关强度'],
    ['交叉载荷差值', '≥ 0.20（主因子 – 次因子）', '删除该题项', '防止题项归属模糊'],
    ['KMO取样充分性', '≥ 0.70（可接受）', '若 < 0.60 审查整体量表设计', '矩阵球型性检验前提'],
    ['Bartlett球型检验', 'p < 0.05', '数据不适合EFA，检查量表', '检验相关矩阵非单位阵'],
    ['因子特征根', '≥ 1.0（Kaiser准则）', '参考碎石图（Scree Plot）', '决定保留因子数'],
    ['累计方差解释率', '≥ 50%', '增加因子数或精简量表', '因子结构有效性'],
    ['构念量表题项数', '至少保留 2 题（推荐 3 题）', '题项 < 2 则剔除该构念', '保证测量稳健性'],
]
paras_t4 = find_para('EFA题项与变量筛选标准')
if paras_t4:
    t4 = make_table(t4_headers, t4_rows, header_shade='E76F51', cw=[1.3,1.6,1.7,1.5])
    insert_tbl_after_para(paras_t4[0], t4)
    print('✓ 表4 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表5  信效度检验指标体系（在第二章六-定量研究设计后）
# ══════════════════════════════════════════════════════════════════════════════
t5_headers = ['检验类型', '指标名称', '计算/判断方式', '参考标准', '判定依据']
t5_rows = [
    ['内部一致性信度', 'Cronbach\'s α', 'α=k/(k-1)×(1-Σσi²/σt²)', '≥ 0.70', 'Nunnally(1978)'],
    ['组合信度', 'CR（Composite Reliability）', 'CR=(Σλ)²/[(Σλ)²+Σ(1-λ²)]', '≥ 0.70', 'Hair et al.(2019)'],
    ['收敛效度', 'AVE（平均方差提取量）', 'AVE=mean(λ²)', '≥ 0.50', 'Fornell & Larcker(1981)'],
    ['收敛效度', '标准化因子载荷', 'CFA路径系数（标准化）', '≥ 0.50，p<0.05', 'Anderson & Gerbing(1988)'],
    ['区分效度', 'Fornell-Larcker准则', '√AVE > 最大两两构念相关系数', '各构念√AVE均高于相关系数', 'Fornell & Larcker(1981)'],
    ['结构效度', 'KMO取样充分性', '矩阵行列式与相关矩阵行列式之比', '≥ 0.70', 'Kaiser(1974)'],
    ['结构效度', 'Bartlett球型检验', 'χ²检验相关矩阵非单位阵', 'p < 0.05', 'Bartlett(1954)'],
]
paras_t5 = find_para('信效度检验指标体系')
if paras_t5:
    t5 = make_table(t5_headers, t5_rows, header_shade='2E4057', cw=[1.0,1.4,1.7,1.2,1.4])
    insert_tbl_after_para(paras_t5[0], t5)
    print('✓ 表5 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表6  各层次多阶段抽样方案
# ══════════════════════════════════════════════════════════════════════════════
t6_headers = ['层次划分', '地域层', '渠道层', '具体抽样方式', '目标群体']
t6_rows = [
    ['层次A', '天津本地\n（占比约30%）', '线上', '微博超话/粉丝群整群抽样\n+ 配额控制（居住地=天津）', 'Z世代天津常住居民\n演唱会粉丝'],
    ['层次B', '天津本地', '线下', '演唱会散场后现场便利抽样\n（配额保证各年龄段）', '现场观演人群\n（大张伟等场次拦访）'],
    ['层次C', '外地来津\n（占比约70%）', '线上', 'PPS整群抽样：按省级人口\n比例在各省粉丝群投放', '外省Z世代粉丝\n有来津观演经历/意愿'],
    ['层次D', '全国\n（潜在消费者）', '线上', '配额抽样：筛选"无演唱会\n经历"Z世代（嵌入筛选题）', '有消费能力但未观演的\nZ世代潜在群体'],
]
paras_t6 = find_para('各层次多阶段抽样方案')
if paras_t6:
    t6 = make_table(t6_headers, t6_rows, header_shade='048A81', cw=[0.7,1.1,0.7,2.1,1.5])
    insert_tbl_after_para(paras_t6[0], t6)
    print('✓ 表6 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表7  各层计划发放量与预期有效样本
# ══════════════════════════════════════════════════════════════════════════════
t7_headers = ['层次', '渠道', '计划发放量\n（份）', '预估有效率', '预期有效\n样本量', '占目标\n比例']
t7_rows = [
    ['天津本地（层A+B）', '线上+线下', '200', '75%/90%', '约158', '~35%'],
    ['外地来津（层C）',   '线上',      '350', '75%',      '约263', '~58%'],
    ['潜在消费者（层D）', '线上',      '80',  '80%',      '约64',  '~14%'],
    ['合计',              '线上+线下', '630', '约79%',    '≥450',  '100%'],
    ['实际回收（清洗后）', '—',        '—',   '—',        'N=300', '最终有效样本'],
]
paras_t7 = find_para('各层计划发放量与预期有效样本')
if paras_t7:
    t7 = make_table(t7_headers, t7_rows, header_shade='6B4E71', cw=[1.4,0.9,0.9,0.9,0.8,0.8])
    insert_tbl_after_para(paras_t7[0], t7)
    print('✓ 表7 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 表8  抽样框构建方案
# ══════════════════════════════════════════════════════════════════════════════
t8_headers = ['渠道类型', '具体平台/地点', '抽样框构建方式', '覆盖目标群体']
t8_rows = [
    ['线上-社交媒体', '微博演唱会超话\n小红书观演笔记\n粉丝应援QQ群', '关键词搜索帖子/用户列表\n通过粉丝群管理员合作\n发起讨论帖嵌入问卷', '全国Z世代演唱\n会粉丝群体'],
    ['线上-票务平台', '大麦网、猫眼演出\n票星球', '评论区问卷引流\n购票用户定向推送', '有购票行为的\n演唱会消费者'],
    ['线上-官方账号', '偶像/乐队官方\n微博、抖音', '官方账号粉丝列表\n转发问卷', 'Z世代核心粉丝\n（深度用户）'],
    ['线下-演唱会现场', '天津奥体中心\n天津体育馆', '散场后出口系统拦截\n（每10人抽1人）', '实地观演人群\n（本地+外地来津）'],
    ['线下-周边场所', '演唱会周边商店\n文创店/周边小摊', '便利抽样', '观演后延伸\n消费行为群体'],
]
paras_t8 = find_para('抽样框构建方案')
if paras_t8:
    t8 = make_table(t8_headers, t8_rows, header_shade='4A90D9', cw=[1.1,1.5,2.0,1.5])
    insert_tbl_after_para(paras_t8[0], t8)
    print('✓ 表8 插入完成')

# ══════════════════════════════════════════════════════════════════════════════
# 更新 fig5 雷达图 图片（替换为双组版本）
# 并更新图5标题文字
# ══════════════════════════════════════════════════════════════════════════════
from docx.oxml.ns import qn as qname

# Update fig5 caption text
for p in doc.paragraphs:
    if '图5  九大构念量表均值雷达图' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '图5  九大构念量表均值双组对比雷达图（左：有无偶像分组；右：高低频观演分组，n=300，Likert 1—5分）'
        print('✓ 图5标题已更新')
        break

# Update radar description text in ch5
for p in doc.paragraphs:
    if '图5雷达图显示，九大构念呈现明显的"高动机—低障碍"极化格局' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = (
            '图5雷达图采用双组对比形式，从有无偶像归属（左图）与观演经历高低频（右图）两个维度，揭示九大构念均值的群体差异。'
            '【有无偶像对比】有偶像组（n=248）在SMI（4.40 vs 2.48，Δ=1.92）、PSR（4.17 vs 1.15，Δ=3.03）、'
            'RSA（4.13 vs 2.27，Δ=1.87）三个维度与无偶像组差异最为显著，印证"偶像归属是社交动机与仪式认同的核心激活因子"；'
            '而EEM（情感体验动机）两组差距较小（4.95 vs 4.64），支持"情感体验是Z世代观演的普适驱动"核心结论。'
            '【高低频观演对比】高频观演组（≥4次，n=104）的TWI（旅游消费意愿，4.18 vs 2.99，Δ=1.19）和PVI（4.60 vs 3.96）'
            '显著高于中低频组，提示观演经历越丰富，旅游延伸消费意愿越强烈。PCB（感知成本障碍）在高频组更低（2.37 vs 3.23），'
            '反映"重复观演者对成本的自我说服与心理适应"效应。'
        )
        print('✓ 图5正文描述已更新')
        break

# Replace the radar image in the document
# Find the image paragraph before/after fig5 caption
for i, p in enumerate(doc.paragraphs):
    if '图5  ' in p.text:
        # The image should be right before this caption paragraph
        # Find image paragraph (inline shape) near this index
        # Check 1-2 paragraphs before
        break

# Actually we need to find the paragraph containing the image and replace it
# The image is inserted as a picture in a paragraph
def replace_image_near_caption(caption_text, new_img_path, img_width_inches=5.8):
    """Find the paragraph containing an image near the caption, replace the image."""
    cap_idx = None
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if caption_text in p.text:
            cap_idx = i
            break
    if cap_idx is None:
        print(f'  Caption not found: {caption_text}')
        return False
    
    # Search in range ±3 paragraphs for an image paragraph
    search_range = range(max(0, cap_idx-3), min(len(paras), cap_idx+3))
    for j in search_range:
        p = paras[j]
        # Check if paragraph contains an inline image
        blips = p._element.findall('.//' + qname('a:blip'), p._element.nsmap)
        if not blips:
            # Also try without namespace
            for elem in p._element.iter():
                if 'blip' in elem.tag:
                    blips = [elem]
                    break
        if blips or p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            # Found image paragraph, clear and re-insert
            from docx.shared import Inches as _Inches
            # Remove all runs from this paragraph
            for run in p.runs:
                run.text = ''
            # Remove inline shapes
            for r_elem in list(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')):
                p._element.remove(r_elem)
            # Add new image
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(new_img_path, width=_Inches(img_width_inches))
            print(f'  图片替换: {caption_text}')
            return True
    print(f'  未找到图片段落 for: {caption_text}')
    return False

# Try replacing fig5
replaced = replace_image_near_caption('图5  ', 'charts/fig5_radar.png', img_width_inches=6.0)
if not replaced:
    print('  尝试替换失败，尝试清除并重新插入...')

# ── 保存 ─────────────────────────────────────────────────────────────────────
doc.save('wfm 部分.docx')
print('\n✓ wfm 部分.docx 已更新（补全 Ch2/Ch3 表格 + 更新双组雷达图）')
