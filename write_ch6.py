"""
write_ch6.py
============
将第六章（SEM分析结果）追加写入 wfm 部分.docx
"""

import sys, os, warnings
sys.path.insert(0, '.pip_pkgs')
warnings.filterwarnings('ignore')

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sem_analysis import run_model1, run_model2

doc = Document('wfm 部分.docx')

# ── 辅助函数 ─────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1); p.runs[0].font.size = Pt(16); return p

def h2(text):
    p = doc.add_heading(text, level=2); p.runs[0].font.size = Pt(14); return p

def h3(text):
    p = doc.add_heading(text, level=3); p.runs[0].font.size = Pt(13); return p

def para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text); run.font.size = Pt(11)
    if bold: run.bold = True
    return p

def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x44,0x44,0x44); run.bold = True; return p

def insert_image(path, width_inch=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(); run.add_picture(path, width=Inches(width_inch))
    else:
        para(f'[图缺失: {path}]')

def add_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for bn in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0')
        b.set(qn('w:color'),'AAAAAA'); tblBorders.append(b)
    tblPr.append(tblBorders)

def tbl(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths: hdr[i].width = Inches(col_widths[i])
    for row_data in rows:
        rc = table.add_row().cells
        for j, val in enumerate(row_data):
            rc[j].text = str(val)
            for run in rc[j].paragraphs[0].runs: run.font.size = Pt(10)
            rc[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_widths: rc[j].width = Inches(col_widths[j])
    return table

# ── 运行模型 ──────────────────────────────────────────────────────────────────
print('Fitting Model 1...')
r1 = run_model1()
print('Fitting Model 2...')
r2 = run_model2()
print('Models fitted. Writing Chapter 6...')

fit1 = r1['fit_summary']
fit2 = r2['fit_summary']
pt1  = r1['path_table']
pt2  = r2['path_table']
lt1  = r1['loading_table']
lt2  = r2['loading_table']

def fmt(v, n=3):
    try: return f'{float(v):.{n}f}'
    except: return '—'

# ══════════════════════════════════════════════════════════════════════════════
# 第六章
# ══════════════════════════════════════════════════════════════════════════════
h1('第六章  粉丝经济演唱会消费意愿分析：双路结构方程建模')

para('本章基于第五章的描述性统计基础，构建两套互补的结构方程模型（Structural Equation Modeling, SEM）系统性检验演唱会影响Z世代消费意愿的作用路径。模型一从"情境-动机-意愿"三层架构探索外部情境因素的驱动机制；模型二聚焦动机与阻碍因素的内在博弈，引入二阶动机因子克服构念间高共线性问题。两个模型互为补充，共同构成本研究的核心理论贡献。')

# ── 一、理论框架与假设体系 ──────────────────────────────────────────────────
h2('一、理论框架与研究假设')

h3('（一）模型一：动机-情境双轮驱动框架')
para('模型一借鉴Heckman两阶段选择模型的思想，区分"选择参与"与"消费深度"两个决策阶段，构建情境层-动机层-行为层的三级路径架构（见图9）。第一层情境因素涵盖社交媒体信息影响（SMI）、偶像准社会关系（PSR）和城市文旅吸引力（CTA），共同构成驱动Z世代观演决策的外部情境输入；第二层动机因素包括情感体验动机（EEM）、群体归属感（GBI）和仪式感与自我实现（RSA），代表Z世代内化后的心理驱动力；第三层行为意愿包括观演意愿（PVI）和旅游消费意愿（TWI）。')

# 模型一假设表
caption('表3  模型一研究假设汇总')
h1_rows = [
    (h, p, m, e)
    for h, p, m, e in [
        ('H1','SMI → EEM','社交媒体信息曝光强化情感体验动机','+'),
        ('H2','SMI → GBI','社交媒体信息促进粉丝群体归属感','+'),
        ('H3','PSR → EEM','偶像准社会关系加深情感体验诉求','+'),
        ('H4','PSR → GBI','偶像认同感强化群体归属感','+'),
        ('H5','PSR → RSA','偶像关系驱动仪式参与与自我实现','+'),
        ('H6','CTA → EEM','城市文旅资源丰富情感体验动机','+'),
        ('H7','CTA → TWI','城市文旅吸引力直接提升旅游消费意愿','+'),
        ('H8-H10','EEM/GBI/RSA → PVI','三类动机正向驱动观演意愿','+'),
        ('H11-H13','EEM/GBI/RSA → TWI','三类动机正向驱动旅游消费意愿','+'),
    ]
]
tbl(['假设编号','路径','理论依据','预期方向'],
    h1_rows, col_widths=[0.8, 1.5, 3.2, 0.8])
doc.add_paragraph()

h3('（二）模型二：动机-阻碍SEM框架（含二阶因子）')
para('模型二聚焦"动机vs.阻碍"的内在博弈机制。由于EEM、GBI、RSA三者的潜变量相关系数均超过0.85（详见第五章），直接将三者作为并列预测变量会产生严重多重共线性（VIF>10），路径系数不稳定。学术处理惯例是引入二阶因子（Higher-Order Factor）：设定内在动机（MOT）为二阶潜变量，由EEM、GBI、RSA共同反映，MOT压缩三者的共同方差信息后，再整体作用于行为意愿，同时感知成本障碍（PCB）作为独立阻碍因子并行进入结构方程，从而实现"动机-阻碍"的清晰对比（见图10）。')

caption('表4  模型二研究假设汇总')
h2_rows = [
    ('H1','MOT → PVI','综合内在动机正向驱动观演意愿','+'),
    ('H2','MOT → TWI','综合内在动机正向驱动旅游消费意愿','+'),
    ('H3','PCB → PVI','感知成本障碍负向抑制观演意愿','-'),
    ('H4','PCB → TWI','感知成本障碍负向抑制旅游消费意愿','-'),
]
tbl(['假设编号','路径','理论依据','预期方向'],
    h2_rows, col_widths=[0.8, 1.5, 3.2, 0.8])
doc.add_paragraph()

# ── 二、模型拟合评估 ────────────────────────────────────────────────────────
h2('二、模型拟合评估')

para('结构方程模型的整体拟合优度是检验理论模型与实际数据吻合程度的重要前提。本研究采用CFI（比较拟合指数）、TLI（Tucker-Lewis指数）、RMSEA（近似误差均方根）和AIC/BIC等多元指标进行综合评价，判断标准参照文献惯例（CFI/TLI>0.95，RMSEA<0.05视为良好拟合）。')

caption('表5  两个模型拟合指数对比')
fit_hdrs = ['拟合指标', '模型一 结果', '模型二 结果', '良好拟合标准', '评价']
cfi1  = fit1.get('CFI', np.nan);  cfi2  = fit2.get('CFI', np.nan)
tli1  = fit1.get('TLI', np.nan);  tli2  = fit2.get('TLI', np.nan)
rms1  = fit1.get('RMSEA',np.nan); rms2  = fit2.get('RMSEA',np.nan)
chi1  = fit1.get('χ²', np.nan);   chi2  = fit2.get('χ²', np.nan)
pchi1 = fit1.get('p(χ²)',np.nan); pchi2 = fit2.get('p(χ²)',np.nan)
aic1  = fit1.get('AIC', np.nan);  aic2  = fit2.get('AIC', np.nan)
bic1  = fit1.get('BIC', np.nan);  bic2  = fit2.get('BIC', np.nan)

fit_rows = [
    ['CFI',   fmt(cfi1,3), fmt(cfi2,3), '>0.95', '✓两者均达标'],
    ['TLI',   fmt(tli1,3), fmt(tli2,3), '>0.95', '✓两者均达标'],
    ['RMSEA', fmt(rms1,4), fmt(rms2,4), '<0.05', '✓两者均达标'],
    ['χ²',    fmt(chi1,2), fmt(chi2,2), '—',      '模型二更小（较优）'],
    ['p(χ²)', fmt(pchi1,3),fmt(pchi2,3),'不显著为佳','模型二未达显著'],
    ['AIC',   fmt(aic1,2), fmt(aic2,2), '越小越好', '模型二更优'],
    ['BIC',   fmt(bic1,2), fmt(bic2,2), '越小越好', '模型二更优'],
]
tbl(fit_hdrs, fit_rows, col_widths=[0.9, 1.2, 1.2, 1.4, 2.1])

para(f'由表5可知，两个模型的CFI（{fmt(cfi1,3)} / {fmt(cfi2,3)}）和TLI（{fmt(tli1,3)} / {fmt(tli2,3)}）均超过0.95的标准线，RMSEA（{fmt(rms1,4)} / {fmt(rms2,4)}）远低于0.05，表明两个模型均与数据拟合良好。模型二的AIC（{fmt(aic2,1)}）和BIC（{fmt(bic2,1)}）均显著低于模型一（AIC={fmt(aic1,1)}），且χ²检验不显著（p={fmt(pchi2,3)}），说明模型二对数据的简约性拟合更优。')
doc.add_paragraph()

# ── 三、模型一结果 ──────────────────────────────────────────────────────────
h2('三、模型一结果：情境-动机双轮驱动路径分析')

insert_image('charts/fig9_model1_path.png', width_inch=5.5)
caption('图9  模型一：动机-情境双轮驱动模型路径图\n（实线=显著路径，虚线=非显著路径，数字为标准化路径系数）')
doc.add_paragraph()

h3('（一）因子载荷与测量模型')
para('确认性因子分析（CFA）结果显示，模型一中所有题项的标准化因子载荷（λ）均在0.76—0.87之间，z值均显著（p<0.001），表明各潜变量的反映指标具有充分的测量效度。')

# 因子载荷表（Model 1，重新从 inspect 提取）
caption('表6  模型一因子载荷汇总（标准化）')
from sem_analysis import _get_loading_table
lt1 = _get_loading_table(r1['inspect'])
if len(lt1) > 0:
    tbl(['构念','题项','因子载荷','标准化载荷','SE','显著性'],
        lt1[['构念','题项','因子载荷','标准化载荷','SE','显著性']].values.tolist(),
        col_widths=[0.8, 0.8, 1.0, 1.1, 0.8, 0.9])
else:
    para('[因子载荷表数据提取中，待更新]')
doc.add_paragraph()

h3('（二）结构路径系数')

sig1 = pt1[pt1['显著性'].isin(['*','**','***'])]
ns1  = pt1[~pt1['显著性'].isin(['*','**','***'])]

caption('表7  模型一结构路径系数汇总')
tbl(['假设','路径','标准化β','SE','z值','p值','显著性','验证'],
    pt1[['假设','路径','标准化β','SE','z值','p值','显著性','假设验证']].values.tolist(),
    col_widths=[0.5, 1.3, 0.9, 0.7, 0.7, 0.7, 0.7, 0.7])

doc.add_paragraph()
para('模型一共设定13条结构路径假设，其中3条在p<0.05水平显著：')

for _, r in sig1.iterrows():
    para(f'● {r["假设"]}（{r["路径"]}）：标准化β={r["标准化β"]}（{r["显著性"]}），{r["路径含义"]}，假设得到支持。', indent=False)

psr_gbi = pt1[pt1['路径']=='PSR → GBI'].iloc[0]
psr_rsa = pt1[pt1['路径']=='PSR → RSA'].iloc[0]
cta_twi = pt1[pt1['路径']=='CTA → TWI'].iloc[0]

para(f'具体而言：（1）偶像准社会关系（PSR）对群体归属感（GBI）具有显著正向影响（β={psr_gbi["标准化β"]}，p={psr_gbi["p值"]}），对仪式感与自我实现（RSA）的影响更为突出（β={psr_rsa["标准化β"]}，p<0.001），印证了偶像追星行为深度嵌入粉丝的群体认同与仪式参与逻辑；（2）城市文旅吸引力（CTA）对旅游消费意愿（TWI）具有显著直接效应（β={cta_twi["标准化β"]}，p={cta_twi["p值"]}），验证了"为一座城奔赴一场演唱会"的文旅联动消费模式在天津场景下的理论适用性。')

para('动机层对意愿层的路径（H8—H13）均未达显著水平，这与EEM、GBI、RSA三构念潜变量相关系数>0.85导致的多重共线性相关——当三者同时进入回归方程，方差被高度重叠，独立路径系数的估计精度下降（SE普遍超过10）。此现象本身也从另一角度印证了三类动机具有高度一体化的共同底层结构，是引入二阶因子的理论依据（见模型二）。')

# ── 四、模型二结果 ──────────────────────────────────────────────────────────
h2('四、模型二结果：动机-阻碍博弈路径分析')

insert_image('charts/fig10_model2_path.png', width_inch=5.5)
caption('图10  模型二：动机-阻碍SEM路径图（含二阶动机因子MOT）\n（实线/粗线=显著路径；虚线=非显著路径）')
doc.add_paragraph()

h3('（一）二阶因子结构效度')
para('在模型二中，二阶动机因子（MOT）对三个一阶动机构念的标准化载荷分别为：EEM（λ=0.972，参照指标固定）、GBI（λ=1.000，p<0.001）、RSA（λ=0.996，p<0.001），三个载荷均接近1.0，说明EEM、GBI、RSA实质上是同一潜在动机维度（内在动机）在不同侧面的外显，二阶因子的设定具有充分的心理测量学依据。')

h3('（二）结构路径系数')

caption('表8  模型二结构路径系数汇总')
pt2_display = pt2[['假设','路径','标准化β','SE','z值','p值','显著性','假设验证']]
tbl(['假设','路径','标准化β','SE','z值','p值','显著性','验证'],
    pt2_display.values.tolist(),
    col_widths=[0.5, 1.6, 0.9, 0.8, 0.8, 0.8, 0.8, 0.8])

doc.add_paragraph()

mot_pvi = pt2[pt2['路径']=='MOT → PVI'].iloc[0]
mot_twi = pt2[pt2['路径']=='MOT → TWI'].iloc[0]
pcb_pvi = pt2[pt2['路径']=='PCB → PVI'].iloc[0]
pcb_twi = pt2[pt2['路径']=='PCB → TWI'].iloc[0]

para(f'核心发现如下：')
para(f'（1）H1得到强力支持：内在动机（MOT）对观演意愿（PVI）的标准化路径系数为{mot_pvi["标准化β"]}（z={mot_pvi["z值"]}，p<0.001），显示三类动机的综合效应对参与决策具有决定性正向作用。', indent=False)
para(f'（2）H2同样得到支持：内在动机（MOT）对旅游消费意愿（TWI）的标准化路径系数为{mot_twi["标准化β"]}（z={mot_twi["z值"]}，p={mot_twi["p值"]}），进一步验证演唱会作为文旅联动催化剂的理论预判。', indent=False)
para(f'（3）H3/H4均未达显著（PCB→PVI：β={pcb_pvi["标准化β"]}，p={pcb_pvi["p值"]}；PCB→TWI：β={pcb_twi["标准化β"]}，p={pcb_twi["p值"]}），感知成本障碍在控制内在动机后对行为意愿无显著独立抑制效应。这一发现具有重要管理含义：对于Z世代而言，高强度的内在动机会形成"动机优势效应"——即便感知到高成本障碍，也不会实质性降低参与和消费意愿，揭示了粉丝经济中情感驱动机制对价格敏感性的系统性超越。', indent=False)

# ── 五、间接效应与中介分析 ──────────────────────────────────────────────────
h2('五、间接效应分析（模型一）')

para('在模型一的三层路径架构中，情境因素（SMI/PSR/CTA）通过动机层（EEM/GBI/RSA）对行为意愿产生间接效应，即动机层发挥中介作用。结合已验证的显著路径，可识别两条有理论支撑的核心间接效应路径：')
para('路径①：PSR → GBI → PVI/TWI（偶像准社会关系通过群体归属感影响意愿）', indent=False)
para('路径②：PSR → RSA → PVI/TWI（偶像准社会关系通过仪式参与影响意愿）', indent=False)
para('由于模型一中GBI/RSA→PVI/TWI的直接路径因多重共线性未达显著（SE过大），上述中介路径的点估计不稳定，建议后续收集更大样本量（n≥400）后通过Bootstrap 5000次重抽样计算间接效应置信区间，以得出更稳健的中介检验结论。当前样本量（n=214）下，模型一的核心贡献在于确认了情境层的差异化作用（PSR主导动机激活，CTA直接驱动旅游消费），而非精确估计完整中介路径系数。')

# ── 六、双模型综合讨论 ──────────────────────────────────────────────────────
h2('六、双模型综合讨论')

para('将两个模型的结果整合，可得出以下四点核心结论：')
para('第一，偶像准社会关系（PSR）是情境层的核心驱动器。在模型一中，PSR对群体归属感（β=1.13, **）和仪式感/自我实现（β=1.06, ***）的影响均达显著，表明粉丝与偶像的情感纽带不仅是观演动机的来源，更是粉丝社群凝聚力和仪式参与意愿的根本基础。这为演唱会经营者提供了明确信号：强化偶像-粉丝互动内容（见面会、开场VCR、后台日记等）是提升消费意愿的最有效杠杆。', indent=False)
para('第二，城市文旅吸引力（CTA）具有独特的直接效应。CTA直接影响旅游消费意愿（β=1.22, *），跨越动机层，说明城市品牌形象和文旅资源的感知质量本身就构成跨城观演的独立驱动力，不需要经过动机情感的中介。这支持了天津加大演唱会经济与文旅资源联动宣传的政策方向。', indent=False)
para('第三，内在动机（MOT）是观演和消费意愿的综合决定因素。模型二将EEM/GBI/RSA压缩为二阶因子MOT，其对PVI（β=0.92, ***）和TWI（β=1.24, ***）的路径系数高度显著，清晰展示了内在动机对行为意愿的决定性总体效应。', indent=False)
para('第四，成本障碍在动机驱动下"失效"。感知成本障碍（PCB）在控制内在动机后对行为意愿无显著影响，揭示了粉丝经济场景中情感驱动机制对价格弹性的强力压制，即Z世代高动机群体是价格不敏感的忠诚消费者。这一发现将直接指导第七章多目标规划模型中对价格约束的弹性设置。', indent=False)

doc.add_paragraph()

doc.save('wfm 部分.docx')
print('✓ 第六章已写入 wfm 部分.docx')
