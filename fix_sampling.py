"""
更新 wfm 部分.docx 中"一、抽样方法"一节（元素索引 60-70）
按用户实际方案重写：地域层×渠道层，A/B/C/D 四层
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy

doc = Document('wfm 部分.docx')
body = doc.element.body
elems = list(body)

# ── 1. 找到 [60] "一、抽样方法" 的元素位置 ─────────────────────────────
def get_text(el):
    return ''.join(t.text or '' for t in el.iter(qn('w:t')))

idx_start = None  # index of "一、抽样方法" heading
idx_end   = None  # index of "二、样本量计算" heading (exclusive)

for i, el in enumerate(elems):
    if el.tag.endswith('}p'):
        t = get_text(el)
        if '一、抽样方法' in t and idx_start is None:
            idx_start = i
        if idx_start and '二、样本量计算' in t:
            idx_end = i
            break

print(f'抽样方法段落: [{idx_start}] → [{idx_end})')

# ── 2. 删除 idx_start+1 .. idx_end-1（保留标题和二的标题本身）──────────
to_remove = elems[idx_start+1 : idx_end]
for el in to_remove:
    body.remove(el)

# 重新获取元素列表，找到"二、样本量计算"位置
elems2 = list(body)
anchor = None
for i, el in enumerate(elems2):
    if el.tag.endswith('}p') and '二、样本量计算' in get_text(el):
        anchor = el
        break

# ── 3. 辅助函数 ──────────────────────────────────────────────────────────
def insert_before(new_el, ref_el):
    """在 ref_el 前插入 new_el"""
    ref_el.addprevious(new_el)

def make_para(text, style=None, bold=False, indent=False, font_size=11):
    p = doc.add_paragraph()
    p._element.getparent().remove(p._element)  # detach from end
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if style:
        p.style = doc.styles[style]
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.24)
    return p._element

def add_border(tbl_obj, color='AAAAAA'):
    tblPr = tbl_obj._tbl.tblPr
    tb = OxmlElement('w:tblBorders')
    for bn in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0');    b.set(qn('w:color'), color)
        tb.append(b)
    tblPr.append(tb)

def set_cell_shade(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), fill_hex); tcPr.append(shd)

def make_table(headers, rows, cw, header_color='2E4057'):
    t = doc.add_table(rows=1, cols=len(headers))
    t._element.getparent().remove(t._element)  # detach
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_border(t)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        cell.width = Inches(cw[i])
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]; run.bold = True
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(255,255,255)
        set_cell_shade(cell, header_color)
    for row in rows:
        rc = t.add_row().cells
        for j, v in enumerate(row):
            rc[j].text = str(v); rc[j].width = Inches(cw[j])
            rc[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if rc[j].paragraphs[0].runs:
                rc[j].paragraphs[0].runs[0].font.size = Pt(10)
    return t._element

# ── 4. 新内容段落 ────────────────────────────────────────────────────────

# 总体描述
p_overview = make_para(
    '本研究采用分层抽样与多阶段抽样相结合的方法，沿"地域层×渠道层"两个维度同时分层，'
    '形成A—D四个独立抽样层次，各层内部采用对应的抽样方式，兼顾样本代表性与现实可操作性。',
    indent=True
)

# （一）地域分层
p_h1 = make_para('（一）地域分层', bold=True)
p_geo = make_para(
    '根据演唱会市场数据，天津大型演唱会本地观众比例约为30%—35%，外地来津观众约占65%—70%。'
    '据此将总样本按30∶70的比例分配至"天津本地层"与"非天津本地层"，'
    '确保外溢消费效应的研究具有足够统计功效。',
    indent=True
)

# （二）渠道分层与各层抽样方式
p_h2 = make_para('（二）渠道分层与各层抽样方式', bold=True)
p_chan = make_para(
    '在两个地域层内部进一步区分线上与线下两个子渠道，共形成A—D四个抽样层。'
    '天津本地层中，线上以市辖区为单位进行整群抽样（层次A），线下在主要商圈实施便利抽样（层次B）；'
    '非天津本地层中，线上以省为单位按人口规模实施PPS整群抽样（层次C），线下在各省进行配额抽样（层次D）。'
    '各层具体抽样方案详见表6。',
    indent=True
)

# 表6 标题
p_cap = make_para('表6  各层次多阶段抽样方案', bold=False)

# 表6 数据
headers = ['层次划分', '地域层', '渠道层', '具体抽样方式', '目标群体']
rows = [
    ['层次A', '天津本地\n（占比约30%）', '线上', '按照市辖区划分整群抽样', 'Z世代天津常住居民'],
    ['层次B', '',                         '线下', '天津主要商圈现场便利抽样', '天津本地Z世代代表人群'],
    ['层次C', '非天津本地\n（占比约70%）','线上', '以省为单位，按省级人口PPS整群抽样', '外省Z世代人群'],
    ['层次D', '',                         '线下', '在各省进行配额抽样', '外省Z世代代表人群'],
]
cw = [0.75, 1.25, 0.7, 2.3, 1.7]
t_el = make_table(headers, rows, cw, header_color='048A81')

# 空段
p_blank = make_para('')

# ── 5. 按顺序插入到 anchor（"二、样本量计算"）之前 ─────────────────────
for el in [p_overview, p_h1, p_geo, p_h2, p_chan, p_cap, t_el, p_blank]:
    anchor.addprevious(el)

# ── 6. 保存 ────────────────────────────────────────────────────────────
doc.save('wfm 部分.docx')
print('✓ 已更新"一、抽样方法"内容，保存至 wfm 部分.docx')
